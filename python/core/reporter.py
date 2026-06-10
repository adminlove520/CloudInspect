# coding: utf-8
"""
多格式报告生成器
"""

import os
import json
from datetime import datetime


class Reporter:
    def __init__(self, config, results):
        self.config = config
        self.results = results
        self.args = config.args

    def generate(self, fmt):
        """生成指定格式报告"""
        output = self.args.get("output") or self.get_default_path(fmt)

        if fmt == "html":
            return self.generate_html(output)
        elif fmt == "json":
            return self.generate_json(output)
        elif fmt == "md":
            return self.generate_md(output)
        elif fmt == "docx":
            return self.generate_docx(output)
        else:
            return self.generate_html(output)

    def get_default_path(self, fmt):
        """获取默认输出路径"""
        import tempfile
        hostname = os.uname().nodename
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        ext = fmt
        return os.path.join("/tmp/cloudinspect", f"inspect_{hostname}_{ts}.{ext}")

    def generate_html(self, path):
        """生成 HTML 报告"""
        os.makedirs(os.path.dirname(path), exist_ok=True)

        results = self.results
        warnings = results.get("warnings", 0)
        critical = results.get("critical", 0)
        os_info = results.get("os_info", {})
        mode = results.get("mode", "routine")

        risk_level = "正常"
        risk_cls = "green"
        if critical > 0:
            risk_level = "严重告警"
            risk_cls = "red"
        elif warnings > 0:
            risk_level = "存在警告"
            risk_cls = "orange"

        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>CloudInspect 云主机安全巡检报告</title>
<style>
  :root {{
    --c-primary: #1976d2; --c-primary-dark: #0d47a1; --c-primary-soft: #e8f1fc;
    --c-text: #1f2937; --c-text-2: #4b5563; --c-bg: #f3f5f9;
    --c-card: #ffffff; --c-border: #d9dee7;
    --c-ok: #16a34a; --c-ok-bg: #f0fdf4;
    --c-warn: #d97706; --c-warn-bg: #fffbeb;
    --c-crit: #dc2626; --c-crit-bg: #fef2f2;
    --c-side-bg: #0f172a; --c-side-text: #cbd5e1; --c-side-muted: #64748b;
    --font-sans: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif;
    --font-mono: ui-monospace, "SFMono-Regular", Consolas, monospace;
  }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: var(--font-sans); font-size: 14px; background: var(--c-bg); color: var(--c-text); line-height: 1.6; }}
  .toc {{ position: fixed; left: 0; top: 0; bottom: 0; width: 220px; background: var(--c-side-bg); padding: 20px 0; overflow-y: auto; z-index: 10; }}
  .toc-brand {{ padding: 0 20px 14px; border-bottom: 1px solid rgba(255,255,255,0.08); margin-bottom: 12px; }}
  .toc-brand .name {{ font-size: 14px; font-weight: 700; color: #fff; }}
  .toc-brand .ver {{ font-size: 11px; color: var(--c-side-muted); font-family: var(--font-mono); }}
  .toc a {{ display: flex; align-items: center; gap: 8px; padding: 8px 20px; color: var(--c-side-text); font-size: 13px; text-decoration: none; border-left: 3px solid transparent; transition: all 0.12s; }}
  .toc a:hover {{ background: rgba(255,255,255,0.05); color: #fff; }}
  .toc a.active {{ background: rgba(25,118,210,0.18); border-left-color: var(--c-primary); color: #fff; }}
  .container {{ max-width: 1200px; margin: 24px auto 24px 240px; padding: 0 24px; }}
  .header {{ background: linear-gradient(135deg, var(--c-primary) 0%, var(--c-primary-dark) 100%); color: #fff; padding: 20px 24px 16px; border-radius: 10px; margin-bottom: 18px; }}
  .header h1 {{ font-size: 20px; font-weight: 700; }}
  .header h1 .tag {{ display: inline-block; background: rgba(255,255,255,0.20); font-size: 11px; padding: 3px 8px; border-radius: 4px; margin-right: 8px; }}
  .header-meta {{ display: flex; flex-wrap: wrap; gap: 16px 28px; margin-top: 10px; font-size: 13px; opacity: 0.9; }}
  .dashboard {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 14px; margin-bottom: 18px; }}
  .dash-card {{ background: var(--c-card); border-radius: 10px; padding: 16px 18px; box-shadow: 0 1px 3px rgba(15,23,42,0.06); border: 1px solid var(--c-border); }}
  .dash-card.green {{ border-left: 4px solid var(--c-ok); }}
  .dash-card.orange {{ border-left: 4px solid var(--c-warn); }}
  .dash-card.red {{ border-left: 4px solid var(--c-crit); }}
  .dash-card.blue {{ border-left: 4px solid var(--c-primary); }}
  .dash-label {{ font-size: 12px; color: var(--c-text-2); margin-bottom: 4px; }}
  .dash-value {{ font-size: 26px; font-weight: 700; }}
  .dash-value.green {{ color: var(--c-ok); }}
  .dash-value.orange {{ color: var(--c-warn); }}
  .dash-value.red {{ color: var(--c-crit); }}
  .dash-sub {{ font-size: 11px; color: var(--c-text-2); margin-top: 2px; }}
  .card {{ background: var(--c-card); border-radius: 10px; margin-bottom: 16px; box-shadow: 0 1px 3px rgba(15,23,42,0.06); border: 1px solid var(--c-border); overflow: hidden; }}
  .card-header {{ padding: 14px 18px; border-bottom: 1px solid var(--c-border); background: var(--c-bg); }}
  .card-header h2 {{ font-size: 15px; font-weight: 600; }}
  .card-body {{ padding: 16px 18px; }}
  .card-body h3 {{ font-size: 13px; font-weight: 600; color: var(--c-text-2); margin: 14px 0 8px; }}
  .metrics-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(160px, 1fr)); gap: 12px; margin-bottom: 14px; }}
  .metric-card {{ background: var(--c-bg); border-radius: 8px; padding: 12px 14px; text-align: center; }}
  .metric-card.green {{ background: var(--c-ok-bg); }}
  .metric-card.orange {{ background: var(--c-warn-bg); }}
  .metric-card.red {{ background: var(--c-crit-bg); }}
  .metric-label {{ font-size: 11px; color: var(--c-text-2); margin-bottom: 4px; }}
  .metric-value {{ font-size: 22px; font-weight: 700; }}
  .metric-card.green .metric-value {{ color: var(--c-ok); }}
  .metric-card.orange .metric-value {{ color: var(--c-warn); }}
  .metric-card.red .metric-value {{ color: var(--c-crit); }}
  .info-table {{ width: 100%; border-collapse: collapse; margin-bottom: 12px; }}
  .info-table td {{ padding: 6px 12px; border-bottom: 1px solid var(--c-border-light); font-size: 13px; }}
  .info-table td:first-child {{ font-weight: 500; color: var(--c-text-2); }}
  pre {{ background: #0f172a; color: #e2e8f0; padding: 12px 14px; border-radius: 6px; font-family: var(--font-mono); font-size: 12px; line-height: 1.5; overflow-x: auto; white-space: pre-wrap; word-break: break-all; }}
  .alert {{ padding: 12px 16px; border-radius: 8px; margin-top: 12px; }}
  .alert-warning {{ background: var(--c-warn-bg); border: 1px solid var(--c-warn); color: #92400e; }}
  .alert-danger {{ background: var(--c-crit-bg); border: 1px solid var(--c-crit); color: #991b1b; }}
  .badge {{ display: inline-block; padding: 2px 7px; border-radius: 10px; font-size: 11px; font-weight: 600; }}
  .badge.ok {{ background: var(--c-ok-bg); color: var(--c-ok); }}
  .badge.warning {{ background: var(--c-warn-bg); color: var(--c-warn); }}
  .badge.critical {{ background: var(--c-crit-bg); color: var(--c-crit); }}
  code {{ background: var(--c-bg); padding: 1px 5px; border-radius: 3px; font-family: var(--font-mono); font-size: 12px; }}
</style>
</head>
<body>

<nav class="toc">
  <div class="toc-brand">
    <div class="name">🔒 CloudInspect</div>
    <div class="ver">v1.0 云主机巡检</div>
  </div>
  <a href="#top" class="active">🏠 总览仪表盘</a>
  <a href="#sysinfo">🖥️ 系统信息</a>
  <a href="#disk">💾 磁盘状态</a>
  <a href="#network">🌐 网络状态</a>
  <a href="#process">💻 进程状态</a>
  <a href="#service">⚙️ 服务状态</a>
  <a href="#security">🔒 安全基线</a>
  <a href="#backdoor">🚨 后门检测</a>
  <a href="#log_analysis">📋 日志分析</a>
  <a href="#history">📝 历史命令</a>
</nav>

<div class="container">
  <div class="header" id="top">
    <h1><span class="tag">v1.0</span> 云主机安全巡检报告</h1>
    <div class="header-meta">
      <span>👤 主机: {os.uname().nodename}</span>
      <span>📅 时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</span>
      <span>🖥️ 操作系统: {os_info.get('pretty', '未知')}</span>
      <span>⏱️ 模式: {mode}</span>
    </div>
  </div>

  <div class="dashboard">
    <div class="dash-card {risk_cls}">
      <div class="dash-label">风险等级</div>
      <div class="dash-value {risk_cls}">{risk_level}</div>
      <div class="dash-sub">严重: {critical} | 警告: {warnings}</div>
    </div>
    <div class="dash-card green">
      <div class="dash-label">CPU</div>
      <div class="dash-value">{results.get('results', {}).get('sysinfo', {{}}).get('cpu_usage', 'N/A')}%</div>
      <div class="dash-sub">{os_info.get('machine', '未知')}</div>
    </div>
    <div class="dash-card green">
      <div class="dash-label">内存</div>
      <div class="dash-value">{results.get('results', {{}}).get('sysinfo', {{}}).get('mem_usage', 'N/A')}%</div>
      <div class="dash-sub">使用率</div>
    </div>
    <div class="dash-card blue">
      <div class="dash-label">操作系统</div>
      <div class="dash-value" style="font-size:16px">{os_info.get('family', 'unknown').upper()}</div>
      <div class="dash-sub">{os_info.get('id', '未知')}</div>
    </div>
  </div>

  <div class="card" id="sysinfo">
    <div class="card-header"><h2>🖥️ 系统信息</h2></div>
    <div class="card-body">
      <table class="info-table">
        <tr><td>主机名</td><td><code>{os.uname().nodename}</code></td></tr>
        <tr><td>操作系统</td><td>{os_info.get('pretty', '未知')}</td></tr>
        <tr><td>内核版本</td><td><code>{os.uname().release}</code></td></tr>
        <tr><td>架构</td><td>{os_info.get('machine', '未知')}</td></tr>
      </table>
    </div>
  </div>

  <div class="card" id="results">
    <div class="card-header"><h2>📊 检测结果汇总</h2></div>
    <div class="card-body">
      <table class="info-table">
        <tr><th>模块</th><th>状态</th></tr>
"""

        # 遍历各模块结果
        for module_name, module_result in results.get("results", {}).items():
            status = module_result.get("status", "unknown")
            badge_cls = "ok" if status == "ok" else ("warning" if status == "skipped" else "critical")
            html += f'        <tr><td>{module_name}</td><td><span class="badge {badge_cls}">{status.upper()}</span></td></tr>\n'

        html += """
      </table>
    </div>
  </div>

  <div style="text-align:center; padding: 20px 0; color: #94a3b8; font-size: 12px; border-top: 1px solid #e2e8f0; margin-top: 24px;">
    <p>CloudInspect v1.0 - 云主机安全巡检报告 | 生成时间: """ + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + """</p>
  </div>
</div>
</body>
</html>"""

        with open(path, "w", encoding="utf-8") as f:
            f.write(html)

        return path

    def generate_json(self, path):
        """生成 JSON 报告"""
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(self.results, f, ensure_ascii=False, indent=2)
        return path

    def generate_md(self, path):
        """生成 Markdown 报告"""
        os.makedirs(os.path.dirname(path), exist_ok=True)
        results = self.results
        md = f"""# CloudInspect 云主机安全巡检报告

## 基本信息
- **主机**: {os.uname().nodename}
- **操作系统**: {results.get('os_info', {}).get('pretty', '未知')}
- **巡检时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- **工作模式**: {results.get('mode', 'routine')}
- **警告**: {results.get('warnings', 0)}
- **严重**: {results.get('critical', 0)}

## 检测模块结果

| 模块 | 状态 |
|------|------|
"""
        for module, result in results.get("results", {}).items():
            status = result.get("status", "unknown")
            md += f"| {module} | {status} |\n"

        md += f"""
---
*由 CloudInspect v1.0 自动生成*
"""
        with open(path, "w", encoding="utf-8") as f:
            f.write(md)
        return path

    def generate_docx(self, path):
        """生成 DOCX 报告"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            os.makedirs(os.path.dirname(path), exist_ok=True)
            doc = Document()

            # 标题
            title = doc.add_heading("CloudInspect 云主机安全巡检报告", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 基本信息
            doc.add_heading("基本信息", 1)
            results = self.results
            os_info = results.get("os_info", {})

            table = doc.add_table(rows=6, cols=2)
            table.style = "Light Grid Accent 1"
            cells = [
                ("主机名", os.uname().nodename),
                ("操作系统", os_info.get("pretty", "未知")),
                ("内核版本", os.uname().release),
                ("架构", os_info.get("machine", "未知")),
                ("工作模式", results.get("mode", "routine")),
                ("巡检时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ]
            for i, (k, v) in enumerate(cells):
                table.rows[i].cells[0].text = k
                table.rows[i].cells[1].text = str(v)

            # 检测结果
            doc.add_heading("检测结果汇总", 1)
            for module, result in results.get("results", {}).items():
                status = result.get("status", "unknown")
                p = doc.add_paragraph()
                p.add_run(f"• {module}: ").bold = True
                p.add_run(status)

            # 保存
            doc.save(path)
            return path

        except ImportError:
            # python-docx 未安装，fallback 到 HTML
            return self.generate_html(path.replace(".docx", ".html"))